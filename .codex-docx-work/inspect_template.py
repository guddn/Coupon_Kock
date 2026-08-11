from pathlib import Path
from zipfile import ZipFile
import sys

from docx import Document
from lxml import etree


SOURCE = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).with_name("reference.docx")
OUTPUT = Path(sys.argv[2]) if len(sys.argv) > 2 else Path(__file__).with_name("template_inventory.txt")


def text_of_cell(cell) -> str:
    return " | ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


lines: list[str] = []
document = Document(SOURCE)
lines.append(f"sections={len(document.sections)}")
lines.append(f"paragraphs={len(document.paragraphs)}")
lines.append(f"tables={len(document.tables)}")
lines.append(f"inline_shapes={len(document.inline_shapes)}")

for index, paragraph in enumerate(document.paragraphs):
    text = paragraph.text.replace("\n", "\\n").strip()
    if text:
        lines.append(f"P{index:04d}\tstyle={paragraph.style.name!r}\t{text}")

for table_index, table in enumerate(document.tables):
    lines.append(f"TABLE {table_index} rows={len(table.rows)} cols={len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        cells = [text_of_cell(cell) for cell in row.cells]
        lines.append(f"  R{row_index:03d}\t" + " || ".join(repr(value) for value in cells))
    if table._tbl.xpath(".//w:drawing"):
        lines.append(f"  DRAWINGS={len(table._tbl.xpath('.//w:drawing'))}")

for section_index, section in enumerate(document.sections):
    for area_name, area in (("header", section.header), ("first_header", section.first_page_header), ("footer", section.footer), ("first_footer", section.first_page_footer)):
        for paragraph_index, paragraph in enumerate(area.paragraphs):
            if paragraph.text.strip():
                lines.append(
                    f"SECTION {section_index} {area_name} P{paragraph_index}: {paragraph.text.strip()}"
                )

with ZipFile(SOURCE) as package:
    document_xml = etree.fromstring(package.read("word/document.xml"))
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    all_text = "\n".join(
        "".join(node.itertext()).strip()
        for node in document_xml.xpath(".//w:p", namespaces=namespaces)
        if "".join(node.itertext()).strip()
    )
    lines.append("OOXML_TEXT_BEGIN")
    lines.append(all_text)
    lines.append("OOXML_TEXT_END")

OUTPUT.write_text("\n".join(lines), encoding="utf-8")
print(OUTPUT)
